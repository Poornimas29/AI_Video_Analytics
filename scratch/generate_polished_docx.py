# scratch/generate_polished_docx.py
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>')
    tcPr.append(shd)

def set_cell_padding(cell, top=140, bottom=140, left=180, right=180):
    """Set inner padding (margins) of a cell in dxa (20 dxa = 1 pt)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    """Apply clean, subtle borders across all cells of a table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def set_callout_borders(table, border_color="1F4E79", fill_hex="F4F6F9"):
    """Set callout box with a prominent left border and soft background."""
    cell = table.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_padding(cell, top=180, bottom=180, left=240, right=200)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def set_row_properties(row, is_header=False):
    """Set row height protection and repeat header across pages."""
    trPr = row._element.get_or_add_trPr()
    cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
    trPr.append(cantSplit)
    if is_header:
        tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trPr.append(tblHeader)

def create_polished_document():
    doc = Document()

    # 1 Inch Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
        # Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("AI Employee Monitoring System | Architecture & Technical Design")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        frun = fp.add_run("Confidential & Proprietary — Automated Attendance & CCTV Analytics")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(128, 128, 128)

    # Color Palette
    PRIMARY_COLOR = RGBColor(31, 78, 121)    # #1F4E79 Deep Navy
    SECONDARY_COLOR = RGBColor(0, 112, 192)  # #0070C0 Blue
    TEXT_COLOR = RGBColor(40, 40, 40)        # #282828 Charcoal Dark
    MUTED_COLOR = RGBColor(90, 90, 90)

    # --- TITLE HEADER ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("System Architecture & Technical Design Document")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    sub_run = sub_p.add_run("AI-Powered Employee Monitoring, Automated Attendance & Safety Analytics (CCTV-Monitorr)")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(12.5)
    sub_run.font.color.rgb = SECONDARY_COLOR
    sub_run.font.bold = True

    # Callout Metadata Box
    callout_tbl = doc.add_table(rows=1, cols=1)
    callout_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout_tbl.autofit = False
    callout_tbl.columns[0].width = Inches(6.5)
    callout_tbl.rows[0].cells[0].width = Inches(6.5)
    set_callout_borders(callout_tbl, border_color="1F4E79", fill_hex="F4F6F9")

    cp = callout_tbl.cell(0, 0).paragraphs[0]
    cp.paragraph_format.line_spacing = 1.2
    cp.paragraph_format.space_after = Pt(0)
    
    def add_meta_item(p, label, val):
        r1 = p.add_run(label + ": ")
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.5)
        r1.font.bold = True
        r1.font.color.rgb = PRIMARY_COLOR
        r2 = p.add_run(val + "\n")
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = TEXT_COLOR

    add_meta_item(cp, "System Version", "2.0 (Production Release)")
    add_meta_item(cp, "Target Environment", "Overhead CCTV Industrial/Office Surveillance Streams")
    add_meta_item(cp, "Core Features", "Touchless Attendance, Continuous Re-ID Tracking, Face ArcFace, Mobile Phone Usage Detection, Live Multi-Camera HUD")
    add_meta_item(cp, "Processing Modes", "Live Multi-Camera RTSP Streams & Fast-Forward Video Playback (1X, 4X, 8X, Uncapped)")
    # Remove trailing newline from last item
    if cp.runs:
        cp.runs[-1].text = cp.runs[-1].text.rstrip("\n")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- HELPERS FOR SECTIONS ---
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return h

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.name = 'Calibri'
            r1.font.size = Pt(10.5)
            r1.font.bold = True
            r1.font.color.rgb = TEXT_COLOR
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = TEXT_COLOR
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3.5)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(bold_prefix + ": ")
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = PRIMARY_COLOR
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10)
        r2.font.color.rgb = TEXT_COLOR
        return p

    # --- SECTION 1 ---
    add_heading_1("1. Problem Statement & Operational Challenges")
    add_body(
        "Standard employee monitoring, attendance logging, and safety compliance methods encounter significant operational limitations when applied to real-world industrial and workplace camera setups:"
    )

    # Clean Problem Table (6.5 inches total width)
    prob_tbl = doc.add_table(rows=6, cols=2)
    prob_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    prob_tbl.autofit = False
    set_table_borders(prob_tbl, color="D3D3D3")
    
    col_widths = [Inches(2.2), Inches(4.3)]
    for row in prob_tbl.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    headers = ["Traditional Challenge", "Operational & Business Impact"]
    set_row_properties(prob_tbl.rows[0], is_header=True)
    for i, h in enumerate(headers):
        cell = prob_tbl.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "1F4E79")
        set_cell_padding(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    prob_data = [
        ("Manual Punch / Biometric Stations", "Creates shift-change queues, risk of proxy punch-ins ('buddy punching'), and fails to capture actual working duration vs. building presence."),
        ("Steep Overhead CCTV Camera Angles", "Ceiling-mounted cameras capture foreshortened bodies, tops of heads, and occluded faces, causing traditional frontal face detectors to fail."),
        ("Uniform Clothing & Worker Ambiguity", "In manufacturing or warehouse bays where all employees wear identical uniforms, standard color trackers and classifiers easily confuse identities."),
        ("Frequent Track Fragmentation", "When employees bend down, turn their backs, or walk behind machines/racks, conventional tracking engines lose identity continuity and spawn duplicate unknown tracks."),
        ("Unauthorized Smartphone Distractions", "Manual safety supervision cannot continuously spot subtle mobile phone usage in safety-sensitive or prohibited factory operational zones.")
    ]

    for row_idx, (c1, c2) in enumerate(prob_data, start=1):
        row = prob_tbl.rows[row_idx]
        set_row_properties(row, is_header=False)
        cell1, cell2 = row.cells[0], row.cells[1]
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_background(cell1, bg)
        set_cell_background(cell2, bg)
        set_cell_padding(cell1, top=100, bottom=100, left=140, right=140)
        set_cell_padding(cell2, top=100, bottom=100, left=140, right=140)

        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(c1)
        r1.font.name = 'Calibri'
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = TEXT_COLOR

        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(c2)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = TEXT_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 2 ---
    add_heading_1("2. Proposed Solution & Architecture Overview")
    add_body(
        "The CCTV-Monitorr platform implements a modern, asynchronous, multi-tier computer vision pipeline designed to convert raw video streams into structured attendance and productivity metrics in real time."
    )
    add_body("The end-to-end processing pipeline is structured across five decoupled tiers:", bold_prefix="Core Architectural Pipeline: ")

    add_bullet("1. Stream Ingestion Tier", "Dedicated reader threads capture RTSP streams via OpenCV FFMPEG with socket timeouts and zero-lag single-slot ring buffering.")
    add_bullet("2. Detection & Tracking Tier", "Ultralytics YOLO (Person, Phone, PPE) accelerated via Apple Metal (MPS) or NVIDIA CUDA paired with ByteTrack trajectory matching.")
    add_bullet("3. Multi-Modal AI Analytics Tier", "Dual identification combining asynchronous InsightFace ArcFace (512-D vectors) with FastReID appearance feature vectors and MediaPipe hand/pose estimation.")
    add_bullet("4. Identity Locking & Attendance Tier", "Consecutive voting lock, continuous working-duration accumulation (excluding absences), and cross-camera session stitching.")
    add_bullet("5. Visualisation & Reporting Tier", "Composite live grid dashboard with HUD overlay, active phone alerts, snapshot capture, and structured daily JSON records.")

    # --- SECTION 3 ---
    add_heading_1("3. Detailed Component Specifications")

    add_heading_2("3.1 Ingestion & Multithreaded Buffering Layer (stream/)")
    add_body(
        "To ensure continuous real-time playback without GUI freezing or latency buildup, the RTSP ingestion pipeline utilizes dedicated reader threads and thread-safe ring buffers."
    )
    add_bullet("RTSPStream", "Connects via OpenCV FFMPEG backend with a 5-second socket timeout (stimeout) to prevent thread locks during network interruptions.")
    add_bullet("FrameBuffer", "Thread-safe single-slot ring buffer protected by mutual exclusion locks. Always yields the freshest available frame, completely eliminating queue latency buildup.")
    add_bullet("Video Timeline Synchronization", "When evaluating pre-recorded video playback (at 1X, 4X, 8X, or uncapped speed), frame timestamps automatically synchronize with the video's internal millisecond offset. This guarantees that attendance logs and working hours remain 100% physically accurate to the recorded event.")

    add_heading_2("3.2 Object Detection & Tracking Layer (detection/ & tracking/)")
    add_body(
        "The detection layer executes a lightweight YOLO singleton detector with multi-class inference for Person (Class 0) and Cell Phone (Class 67)."
    )
    add_bullet("Hardware Auto-Targeting", "Automatically targets Apple Silicon Metal (MPS), NVIDIA CUDA, OpenVINO, or CPU for optimal throughput.")
    add_bullet("ByteTrack Multi-Object Tracker", "Maintains track continuity across frames using Kalman filtering and Hungarian association, assigning stable track IDs to persons.")

    add_heading_2("3.3 Dual Identification: Face Recognition & Body Re-ID (ai/)")
    add_body(
        "To solve the steep overhead CCTV problem where faces are only intermittently visible, the system employs a dual-identification strategy:"
    )
    add_bullet("Asynchronous InsightFace Pool (ArcFace buffalo_l)", "Extracts 512-dimensional normalized facial feature vectors in a 2-worker background thread pool. Frames are filtered for quality using Laplacian blur thresholds and minimum face dimensions (30px).")
    add_bullet("Consecutive Match Identity Locking", "Requires 2 consecutive positive matches with cosine similarity >= 0.50 before permanently locking a Track ID to an employee name, eliminating false-positive switches.")
    add_bullet("FastReID Appearance Recovery", "Extracts L2-normalized deep body embeddings (MobileNetV3 / OSNet). When an employee bends down or walks behind a rack (causing a track ID change), the Re-ID engine searches active sessions (threshold: 0.75) and seamlessly re-links the person to their session.")

    add_heading_2("3.4 Human Pose & Unauthorized Mobile Phone Usage Detection (ai/)")
    add_body(
        "Mobile phone detection combines object detection with pose geometry to eliminate false positives from phones lying on work tables:"
    )
    add_bullet("Keypoint Extraction", "MediaPipe / Heuristic Pose estimator identifies wrists, hands, shoulders, and head orientation.")
    add_bullet("Spatial & Distance Proximity", "Verifies that the phone bounding box is spatially contained within the person's bounding box AND within a 15% distance threshold from the person's hand keypoints.")
    add_bullet("Temporal Confirmation Window", "Enforces continuous usage across a 2.0-second window (PHONE_USAGE_CONFIRM_SECONDS) before flagging a confirmed phone usage event.")

    add_heading_2("3.5 Attendance & Productivity Analytics Engine (session/)")
    add_body(
        "The Attendance and Session Manager tracks real working time and writes structured JSON logs to output/known/ and output/unknown/:"
    )
    add_bullet("Net Continuous Working Duration", "Accurately accumulates time only while the employee is actively tracked on camera, strictly excluding lunch breaks, out-of-frame time, and network disconnects.")
    add_bullet("Session Expiry & Lifecycle", "Holds lost tracks for 60 seconds (TRACK_TIMEOUT) before finalizing the session as 'exited' and archiving attendance.")
    add_bullet("Daily Aggregation & Productivity Score", "Computes Check-In, Check-Out, Net Work Hours, Total Phone Distraction Time, and Productivity Score: Productivity = max(0, 100 * (Work Time - Phone Time) / Work Time).")

    # --- SECTION 4 ---
    add_heading_1("4. Technology Stack & Frameworks")
    add_body("The platform is engineered using industry-standard, production-grade tools and libraries:")

    # Clean Tech Table (6.5 inches total width)
    tech_tbl = doc.add_table(rows=11, cols=3)
    tech_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_tbl.autofit = False
    set_table_borders(tech_tbl, color="D3D3D3")
    
    t_widths = [Inches(1.8), Inches(1.8), Inches(2.9)]
    for row in tech_tbl.rows:
        for i, w in enumerate(t_widths):
            row.cells[i].width = w

    tech_headers = ["Layer / Domain", "Technology & Version", "Role & Engineering Rationale"]
    set_row_properties(tech_tbl.rows[0], is_header=True)
    for i, h in enumerate(tech_headers):
        cell = tech_tbl.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "1F4E79")
        set_cell_padding(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    tech_data = [
        ("Runtime & Language", "Python 3.9+ / Virtual Env", "Core application execution environment with future type annotations."),
        ("Video I/O & Graphics", "OpenCV 4.10+ (FFMPEG)", "Multithreaded RTSP decoding, buffer management, and live HUD composition."),
        ("Object Detection", "Ultralytics YOLO (PyTorch)", "High-accuracy, real-time person & cell-phone detection with Apple Metal (MPS) / CUDA."),
        ("Multi-Object Tracking", "ByteTrack + Kalman Filter", "Trajectory association preserving continuous identities across occlusions."),
        ("Face Recognition", "InsightFace (ArcFace buffalo_l)", "Deep 512-D facial feature embeddings with RetinaFace detector."),
        ("Model Backend", "ONNX Runtime (1.19+)", "High-performance inference engine supporting CPU, CoreML, and CUDA execution providers."),
        ("Body Re-ID", "MobileNetV3 / OSNet (PyTorch)", "L2-normalized full-body appearance embeddings for track re-linking."),
        ("Pose Estimation", "MediaPipe Tasks & Heuristic", "Skeletal keypoints, head pose orientation, and hand-to-phone proximity measurement."),
        ("Concurrency & Async", "Python threading, queue, Lock", "Decoupled async worker pool, single-slot ring buffering, and thread synchronization."),
        ("Data Persistence", "JSON & Structured File System", "Lightweight, zero-dependency storage for daily attendance records and event summaries.")
    ]

    for row_idx, (c1, c2, c3) in enumerate(tech_data, start=1):
        row = tech_tbl.rows[row_idx]
        set_row_properties(row, is_header=False)
        cell1, cell2, cell3 = row.cells[0], row.cells[1], row.cells[2]
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell3.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_background(cell1, bg)
        set_cell_background(cell2, bg)
        set_cell_background(cell3, bg)
        set_cell_padding(cell1, top=90, bottom=90, left=130, right=130)
        set_cell_padding(cell2, top=90, bottom=90, left=130, right=130)
        set_cell_padding(cell3, top=90, bottom=90, left=130, right=130)

        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(c1)
        r1.font.name = 'Calibri'
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = TEXT_COLOR

        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(c2)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9)
        r2.font.color.rgb = TEXT_COLOR

        p3 = cell3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r3 = p3.add_run(c3)
        r3.font.name = 'Calibri'
        r3.font.size = Pt(9)
        r3.font.color.rgb = TEXT_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 5 ---
    add_heading_1("5. Key Engineering Solutions & Production Optimizations")
    add_bullet("1. Asynchronous Worker Race Condition Fix", "drain_results() is executed prior to process_timeouts() to ensure async recognition results arriving just as a track leaves the frame are never discarded.")
    add_bullet("2. Zero-Lock Hot Path Optimization", "Removed synchronous print() calls from frame-processing loops, replacing them with asynchronous logger levels, raising streaming FPS significantly.")
    add_bullet("3. Single-Pass Re-ID Feature Extraction", "Appearance features are extracted exactly once per track per frame, cached in a frame dictionary, and reused across identity and tracking modules.")
    add_bullet("4. Accelerated Video Playback (4X / 8X / Uncapped)", "Enables batch-speed processing of pre-recorded surveillance videos while maintaining 100% video-timeline timestamp accuracy for all generated attendance reports.")
    add_bullet("5. Robust macOS & Cross-Platform Support", "Native Apple Silicon Metal acceleration (MPS) and graceful fallback mechanisms for high-performance deployment on macOS, Linux, and Windows.")

    # Save
    workspace_path = "/Users/poornima/Video Analytics/System_Architecture_and_Design_Document.docx"
    downloads_path = "/Users/poornima/Downloads/System_Architecture_and_Design_Document.docx"
    
    doc.save(workspace_path)
    doc.save(downloads_path)
    print(f"Successfully generated Polished Word document at:\n1. {workspace_path}\n2. {downloads_path}")

if __name__ == "__main__":
    create_polished_document()
