# scratch/generate_word_doc.py
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding in dxa units (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette
    PRIMARY_COLOR = RGBColor(31, 78, 121)    # #1F4E79 Deep Navy
    SECONDARY_COLOR = RGBColor(89, 89, 89)   # Dark Gray
    ACCENT_COLOR = RGBColor(0, 112, 192)     # #0070C0 Blue
    TEXT_COLOR = RGBColor(38, 38, 38)        # Charcoal

    # Document Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("System Architecture & Technical Design Document")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(20)
    run_sub = sub_p.add_run("AI Employee Monitoring & Automated Attendance System (CCTV-Monitorr)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = ACCENT_COLOR
    run_sub.font.italic = True

    # Metadata callout box
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_cell = meta_table.cell(0, 0)
    set_cell_background(meta_cell, "F2F4F7")
    set_cell_margins(meta_cell, top=140, bottom=140, left=200, right=200)
    
    mp = meta_cell.paragraphs[0]
    mp.paragraph_format.space_after = Pt(2)
    r1 = mp.add_run("Document Version: ")
    r1.font.bold = True
    mp.add_run("2.0 (Production Release)\n")
    r2 = mp.add_run("Target Environment: ")
    r2.font.bold = True
    mp.add_run("Overhead CCTV Streams / High-Throughput Edge Analytics\n")
    r3 = mp.add_run("Core Capabilities: ")
    r3.font.bold = True
    mp.add_run("Touchless Attendance, Continuous Body Tracking, Face ArcFace Recognition, Mobile Phone Usage Detection, Multi-Camera Grid Dashboard")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for Headings
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = ACCENT_COLOR
        return h

    def add_body_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = TEXT_COLOR
        r_text = p.add_run(text)
        r_text.font.name = 'Calibri'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = TEXT_COLOR
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r_bold = p.add_run(bold_prefix + ": ")
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
        r_bold.font.color.rgb = TEXT_COLOR
        r_text = p.add_run(text)
        r_text.font.name = 'Calibri'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = TEXT_COLOR
        return p

    # --- SECTION 1 ---
    add_heading_1("1. Problem Statement & Operational Challenges")
    add_body_p(
        "Traditional employee timekeeping, attendance tracking, and factory floor monitoring face critical operational challenges when deployed in practical, real-world environments:"
    )

    # Problem Table
    prob_table = doc.add_table(rows=6, cols=2)
    prob_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Traditional Challenge", "Operational & Business Impact"]
    for i, h in enumerate(headers):
        cell = prob_table.cell(0, i)
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10.5)

    prob_data = [
        ("Manual Punch / Biometric Stations", "Creates shift-change bottlenecks, proxy clock-ins ('buddy punching'), hygiene concerns, and fails to measure actual working duration vs. building presence."),
        ("Steep Overhead CCTV Camera Angles", "Ceiling-mounted cameras capture foreshortened bodies, tops of heads, and occluded faces, causing conventional frontal face detectors to fail."),
        ("Uniform Clothing & Worker Ambiguity", "In manufacturing or warehouse bays where all employees wear identical uniforms, standard color trackers and classifiers easily confuse identities."),
        ("Frequent Track Fragmentation", "When employees bend down, turn their backs, or walk behind machines/racks, conventional tracking engines lose identity continuity and spawn duplicate unknown tracks."),
        ("Unauthorized Smartphone Distractions", "Manual safety supervision cannot continuously spot subtle mobile phone usage in safety-sensitive or prohibited factory operational zones.")
    ]

    for row_idx, (c1, c2) in enumerate(prob_data, start=1):
        cell1 = prob_table.cell(row_idx, 0)
        cell2 = prob_table.cell(row_idx, 1)
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_background(cell1, bg)
        set_cell_background(cell2, bg)
        set_cell_margins(cell1, top=100, bottom=100, left=140, right=140)
        set_cell_margins(cell2, top=100, bottom=100, left=140, right=140)
        
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(c1)
        r1.font.bold = True
        r1.font.size = Pt(10)
        
        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(c2)
        r2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- SECTION 2 ---
    add_heading_1("2. Proposed Solution & Architecture Overview")
    add_body_p(
        "The CCTV-Monitorr platform implements a modern, asynchronous, multi-tier computer vision pipeline that automatically converts raw RTSP camera streams into real-time attendance, tracking, and compliance intelligence."
    )

    add_body_p(
        "The architecture is decoupled into five distinct processing tiers:",
        bold_prefix="Core Architectural Tiers: "
    )
    add_bullet("1. High-Throughput Ingestion Tier", "Non-blocking multithreaded RTSP stream capture with thread-safe single-slot ring buffering and auto-reconnection.")
    add_bullet("2. Spatial Detection & Tracking Tier", "Ultralytics YOLO (Person, Phone, PPE) accelerated via Apple Silicon Metal (MPS) or NVIDIA CUDA paired with ByteTrack trajectory matching.")
    add_bullet("3. Multi-Modal AI Analytics Tier", "Dual identification pipeline combining asynchronous InsightFace ArcFace (512-D vectors) with FastReID appearance feature embeddings and MediaPipe hand/pose estimation.")
    add_bullet("4. Identity Locking & Attendance Tier", "Consecutive voting lock, continuous working-duration accumulation (excluding absences), and cross-camera session stitching.")
    add_bullet("5. Visualisation & Reporting Tier", "Composite live grid dashboard with HUD overlay, active phone alerts, snapshot capture, and structured daily JSON records.")

    # --- SECTION 3 ---
    add_heading_1("3. Detailed Component Specifications")

    add_heading_2("3.1 Ingestion & Multithreaded Buffering Layer (stream/)")
    add_body_p(
        "To ensure continuous real-time playback without GUI freezing or latency buildup, the RTSP ingestion pipeline utilizes dedicated reader threads and thread-safe ring buffers."
    )
    add_bullet("RTSPStream", "Connects to cameras via OpenCV FFMPEG backend with a 5-second socket timeout (stimeout) to avoid thread blocking during network disruptions.")
    add_bullet("FrameBuffer", "Thread-safe single-slot ring buffer protected by mutual exclusion locks. Always yields the freshest available frame, preventing queue latency buildup.")
    add_bullet("Physical Video Timeline Synchronization", "When evaluating pre-recorded video playback (at 1X, 4X, 8X, or uncapped speed), frame timestamps automatically synchronize with the video's internal millisecond offset. This guarantees that attendance logs and working hours remain 100% physically accurate to the recorded event.")

    add_heading_2("3.2 Object Detection & Tracking Layer (detection/ & tracking/)")
    add_body_p(
        "The detection layer executes a lightweight YOLO singleton detector with multi-class inference for Person (Class 0) and Cell Phone (Class 67)."
    )
    add_bullet("Hardware Auto-Targeting", "Automatically selects Apple Silicon Metal (MPS), NVIDIA CUDA, OpenVINO, or CPU to ensure high inference throughput.")
    add_bullet("ByteTrack Multi-Object Tracker", "Maintains track continuity across frames using Kalman filtering and Hungarian association, assigning stable track IDs to persons.")

    add_heading_2("3.3 Dual Identification: Face Recognition & Body Re-ID (ai/)")
    add_body_p(
        "To solve the steep overhead CCTV problem where faces are only intermittently visible, the system employs a dual-identification strategy:"
    )
    add_bullet("Asynchronous InsightFace Pool (ArcFace buffalo_l)", "Extracts 512-dimensional normalized facial feature vectors in a 2-worker background thread pool. Frames are filtered for quality using Laplacian blur thresholds and minimum face dimensions (30px).")
    add_bullet("Consecutive Match Identity Locking", "Requires 2 consecutive positive matches with cosine similarity >= 0.50 before permanently locking a Track ID to an employee name, eliminating false-positive switches.")
    add_bullet("FastReID Appearance Recovery", "Extracts L2-normalized deep body embeddings (MobileNetV3 / OSNet). When an employee bends down or walks behind a rack (causing a track ID change), the Re-ID engine searches active sessions (threshold: 0.75) and seamlessly re-links the person to their session.")

    add_heading_2("3.4 Human Pose & Unauthorized Mobile Phone Usage Detection (ai/)")
    add_body_p(
        "Mobile phone detection combines object detection with pose geometry to eliminate false positives from phones lying on work tables:"
    )
    add_bullet("Keypoint Extraction", "MediaPipe / Heuristic Pose estimator identifies wrists, hands, shoulders, and head orientation.")
    add_bullet("Spatial & Distance Proximity", "Verifies that the phone bounding box is spatially contained within the person's bounding box AND within a 15% distance threshold from the person's hand keypoints.")
    add_bullet("Temporal Confirmation Window", "Enforces continuous usage across a 2.0-second window (PHONE_USAGE_CONFIRM_SECONDS) before flagging a confirmed phone usage event.")

    add_heading_2("3.5 Attendance & Productivity Analytics Engine (session/)")
    add_body_p(
        "The Attendance and Session Manager tracks real working time and writes structured JSON logs to output/known/ and output/unknown/:"
    )
    add_bullet("Net Continuous Working Duration", "Accurately accumulates time only while the employee is actively tracked on camera, strictly excluding lunch breaks, out-of-frame time, and network disconnects.")
    add_bullet("Session Expiry & Lifecycle", "Holds lost tracks for 60 seconds (TRACK_TIMEOUT) before finalizing the session as 'exited' and archiving attendance.")
    add_bullet("Daily Aggregation & Productivity Score", "Computes Check-In, Check-Out, Net Work Hours, Total Phone Distraction Time, and Productivity Score: Productivity = max(0, 100 * (Work Time - Phone Time) / Work Time).")

    # --- SECTION 4 ---
    add_heading_1("4. Technology Stack & Frameworks")
    add_body_p("The system is engineered using industry-standard, production-grade tools and libraries:")

    tech_table = doc.add_table(rows=11, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_headers = ["Layer / Domain", "Technology & Version", "Role & Engineering Rationale"]
    for i, h in enumerate(tech_headers):
        cell = tech_table.cell(0, i)
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    tech_data = [
        ("Runtime & Language", "Python 3.9+ / Virtual Env", "Core application execution environment with future type annotations."),
        ("Video I/O & Graphics", "OpenCV 4.10+ (FFMPEG)", "Multithreaded RTSP decoding, buffer management, and live HUD composition."),
        ("Object Detection", "Ultralytics YOLO (PyTorch)", "High-accuracy, real-time person & cell-phone detection with Apple Metal (MPS) / CUDA."),
        ("Multi-Object Tracking", "ByteTrack + Kalman Filter", "Trajectory association preserving continuous identities across occlusions."),
        ("Face Recognition", "InsightFace (ArcFace buffalo_l)", "Deep 512-D facial feature embeddings with RetinaFace detector."),
        ("Model Execution Backend", "ONNX Runtime (1.19+)", "High-performance inference engine supporting CPU, CoreML, and CUDA execution providers."),
        ("Body Re-Identification", "MobileNetV3 / OSNet (PyTorch)", "L2-normalized full-body appearance embeddings for track re-linking."),
        ("Pose Estimation", "MediaPipe Tasks & Heuristic", "Skeletal keypoints, head pose orientation, and hand-to-phone proximity measurement."),
        ("Concurrency & Async", "Python threading, queue, Lock", "Decoupled async worker pool, single-slot ring buffering, and thread synchronization."),
        ("Data Persistence", "JSON & Structured File System", "Lightweight, zero-dependency storage for daily attendance records and event summaries.")
    ]

    for row_idx, (c1, c2, c3) in enumerate(tech_data, start=1):
        cell1 = tech_table.cell(row_idx, 0)
        cell2 = tech_table.cell(row_idx, 1)
        cell3 = tech_table.cell(row_idx, 2)
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_background(cell1, bg)
        set_cell_background(cell2, bg)
        set_cell_background(cell3, bg)
        set_cell_margins(cell1, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell2, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell3, top=80, bottom=80, left=120, right=120)
        
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(c1)
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        
        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(c2)
        r2.font.size = Pt(9.5)

        p3 = cell3.paragraphs[0]
        r3 = p3.add_run(c3)
        r3.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- SECTION 5 ---
    add_heading_1("5. Key Engineering Solutions & Production Optimizations")
    add_bullet("1. Asynchronous Worker Race Condition Fix", "drain_results() is executed prior to process_timeouts() to ensure async recognition results arriving just as a track leaves the frame are never discarded.")
    add_bullet("2. Zero-Lock Hot Path Optimization", "Removed synchronous print() calls from frame-processing loops, replacing them with asynchronous logger levels, raising streaming FPS significantly.")
    add_bullet("3. Single-Pass Re-ID Feature Extraction", "Appearance features are extracted exactly once per track per frame, cached in a frame dictionary, and reused across identity and tracking modules.")
    add_bullet("4. Accelerated Video Playback (4X / 8X / Uncapped)", "Enables batch-speed processing of pre-recorded surveillance videos while maintaining 100% video-timeline timestamp accuracy for all generated attendance reports.")
    add_bullet("5. Robust macOS & Cross-Platform Support", "Native Apple Silicon Metal acceleration (MPS) and graceful fallback mechanisms for high-performance deployment on macOS, Linux, and Windows.")

    # Save documents
    workspace_path = "/Users/poornima/Video Analytics/System_Architecture_and_Design_Document.docx"
    downloads_path = "/Users/poornima/Downloads/System_Architecture_and_Design_Document.docx"
    
    doc.save(workspace_path)
    doc.save(downloads_path)
    print(f"Successfully generated Word document at:\n1. {workspace_path}\n2. {downloads_path}")

if __name__ == "__main__":
    create_document()
