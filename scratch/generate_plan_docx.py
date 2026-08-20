# scratch/generate_plan_docx.py
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>')
    tcPr.append(shd)

def set_cell_padding(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def set_callout_borders(table, border_color="1F4E79", fill_hex="F4F6F9"):
    cell = table.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_padding(cell, top=180, bottom=180, left=240, right=200)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def set_row_properties(row, is_header=False):
    trPr = row._element.get_or_add_trPr()
    cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
    trPr.append(cantSplit)
    if is_header:
        tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trPr.append(tblHeader)

def create_plan_document():
    doc = Document()

    # 1.0 Inch Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
        # Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Project Implementation & Execution Plan | AI Employee Monitoring")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        frun = fp.add_run("Confidential & Proprietary — CCTV-Monitorr Implementation Roadmap")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(128, 128, 128)

    # Color Palette
    PRIMARY_COLOR = RGBColor(31, 78, 121)    # Deep Navy
    SECONDARY_COLOR = RGBColor(0, 112, 192)  # Bright Blue
    TEXT_COLOR = RGBColor(40, 40, 40)        # Dark Charcoal

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("Project Implementation & Execution Plan")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    sub_run = sub_p.add_run("AI-Powered Touchless Attendance, Re-ID Tracking & Productivity Analytics")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(12.5)
    sub_run.font.bold = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    # Plan Summary Callout Box
    callout_tbl = doc.add_table(rows=1, cols=1)
    callout_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout_tbl.autofit = False
    callout_tbl.columns[0].width = Inches(6.5)
    callout_tbl.rows[0].cells[0].width = Inches(6.5)
    set_callout_borders(callout_tbl, border_color="1F4E79", fill_hex="F4F6F9")

    cp = callout_tbl.cell(0, 0).paragraphs[0]
    cp.paragraph_format.line_spacing = 1.2
    
    def add_meta_item(p, label, val):
        r1 = p.add_run(label + ": ")
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.5)
        r1.font.bold = True
        r1.font.color.rgb = PRIMARY_COLOR
        r2 = p.add_run(val + "\n")
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = TEXT_COLOR

    add_meta_item(cp, "Project Scope", "End-to-end automated CCTV video analytics for employee attendance, real-time Re-ID tracking, and mobile distraction detection.")
    add_meta_item(cp, "Target Deployment", "Edge Workstation / Server connected to Industrial & Workplace Dahua/RTSP Cameras.")
    add_meta_item(cp, "Key Execution Strategy", "Multi-Phase Modular Architecture with Asynchronous Multi-Worker Processing.")
    if cp.runs:
        cp.runs[-1].text = cp.runs[-1].text.rstrip("\n")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Helpers
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return h

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.name = 'Calibri'
            r1.font.size = Pt(10.5)
            r1.font.bold = True
            r1.font.color.rgb = TEXT_COLOR
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = TEXT_COLOR
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3.5)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(bold_prefix + ": ")
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = PRIMARY_COLOR
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10)
        r2.font.color.rgb = TEXT_COLOR
        return p

    # --- SECTION 1: EXECUTION PHASES ---
    add_heading_1("1. Phased Project Execution Roadmap")
    add_body("The project is executed across seven structured phases to ensure stability, high throughput, and fault tolerance:")

    # Table of Phases (6.5 in width)
    phase_tbl = doc.add_table(rows=8, cols=3)
    phase_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    phase_tbl.autofit = False
    set_table_borders(phase_tbl, color="D3D3D3")
    
    col_w = [Inches(1.2), Inches(2.3), Inches(3.0)]
    for row in phase_tbl.rows:
        for i, w in enumerate(col_w):
            row.cells[i].width = w

    headers = ["Phase", "Milestone / Module", "Key Objectives & Deliverables"]
    set_row_properties(phase_tbl.rows[0], is_header=True)
    for i, h in enumerate(headers):
        cell = phase_tbl.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "1F4E79")
        set_cell_padding(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    phase_data = [
        ("Phase 1", "Stream Ingestion & Buffering", "Multithreaded RTSP reader, OpenCV FFMPEG socket timeout, thread-safe ring buffer, auto-reconnection, and mock simulation mode."),
        ("Phase 2", "Detection & Tracking", "Ultralytics YOLO (Person/Phone detection) with hardware auto-acceleration (Metal/CUDA) and ByteTrack trajectory matching."),
        ("Phase 3", "Dual-Identity Recognition", "Async InsightFace ArcFace (512-D) pool + FastReID (MobileNetV3) appearance embeddings for track recovery across occlusions."),
        ("Phase 4", "Pose & Phone Usage Analytics", "MediaPipe / Heuristic pose estimation, hand tracking, spatial bounding-box containment, and temporal 2.0s confirmation window."),
        ("Phase 5", "Attendance & Session Engine", "Consecutive voting identity locking, net continuous working time accumulation, daily JSON attendance generation (known/unknown)."),
        ("Phase 6", "Visualization & Live HUD", "Dynamic multi-camera grid compositor, real-time FPS/status HUD, flashing violation alerts, and snapshot capture."),
        ("Phase 7", "Testing & Deployment Tuning", "Comprehensive 78-test unit suite, 4X/8X accelerated video playback synchronization, and production containerization.")
    ]

    for row_idx, (c1, c2, c3) in enumerate(phase_data, start=1):
        row = phase_tbl.rows[row_idx]
        set_row_properties(row, is_header=False)
        cell1, cell2, cell3 = row.cells[0], row.cells[1], row.cells[2]
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell3.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_background(cell1, bg)
        set_cell_background(cell2, bg)
        set_cell_background(cell3, bg)
        set_cell_padding(cell1, top=90, bottom=90, left=120, right=120)
        set_cell_padding(cell2, top=90, bottom=90, left=120, right=120)
        set_cell_padding(cell3, top=90, bottom=90, left=120, right=120)

        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(c1)
        r1.font.name = 'Calibri'
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = TEXT_COLOR

        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(c2)
        r2.font.name = 'Calibri'
        r2.font.bold = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = TEXT_COLOR

        p3 = cell3.paragraphs[0]
        r3 = p3.add_run(c3)
        r3.font.name = 'Calibri'
        r3.font.size = Pt(9)
        r3.font.color.rgb = TEXT_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 2: HARDWARE & RESOURCE PLAN ---
    add_heading_1("2. Hardware, Environment & Resource Requirements")
    add_body("To ensure consistent 25+ FPS real-time processing across multiple 1080p cameras, the deployment environment adheres to the following specifications:")

    add_bullet("Compute Hardware", "Apple Silicon (M1/M2/M3 with MPS Metal acceleration) or NVIDIA GPU (RTX 3060+ / T4 / A10 with CUDA 11.8+ / 12.x). CPU-only fallback supported via OpenVINO.")
    add_bullet("Memory & Storage", "Minimum 16 GB RAM (32 GB recommended for >=4 camera streams) and high-speed NVMe SSD storage for frame buffering and JSON record persistence.")
    add_bullet("Network Infrastructure", "Gigabit Ethernet LAN dedicated to CCTV subnets with PoE (Power over Ethernet) switches for stable RTSP stream transmission.")
    add_bullet("Camera Specifications", "Standard Dahua / Hikvision / ONVIF IP Cameras streaming H.264 / H.265 at 1080p @ 25 FPS with main-stream resolution for recognition and sub-stream for preview.")

    # --- SECTION 3: RISK MANAGEMENT ---
    add_heading_1("3. Risk Management & Technical Mitigations")

    risk_tbl = doc.add_table(rows=5, cols=3)
    risk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    risk_tbl.autofit = False
    set_table_borders(risk_tbl, color="D3D3D3")
    
    r_widths = [Inches(1.8), Inches(2.2), Inches(2.5)]
    for row in risk_tbl.rows:
        for i, w in enumerate(r_widths):
            row.cells[i].width = w

    r_headers = ["Identified Risk", "Operational Impact", "Implemented Mitigation Strategy"]
    set_row_properties(risk_tbl.rows[0], is_header=True)
    for i, h in enumerate(r_headers):
        cell = risk_tbl.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "1F4E79")
        set_cell_padding(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    risk_data = [
        ("CCTV Network Drops / RTSP Lag", "Freezes UI dashboard and causes decoding queues to build up.", "Dedicated background reader thread, 5s socket timeout, and single-slot ring buffer dropping stale frames."),
        ("Overhead Camera Face Occlusion", "Employees turn backs or look down, causing face recognition to fail.", "FastReID full-body appearance vector extraction to seamlessly re-link active sessions across lost tracks."),
        ("False Identity Switches", "Single noisy frame assigns identity to wrong worker.", "2-Stage Consecutive Match Voting (>=2 matches at cosine similarity >= 0.50) before identity is locked."),
        ("Heavy Model Latency on Frame Rate", "Running ArcFace on 25 FPS video drops stream rate to <5 FPS.", "Decoupled 2-worker asynchronous recognition pool running face embeddings off the main display loop.")
    ]

    for row_idx, (c1, c2, c3) in enumerate(risk_data, start=1):
        row = risk_tbl.rows[row_idx]
        set_row_properties(row, is_header=False)
        cell1, cell2, cell3 = row.cells[0], row.cells[1], row.cells[2]
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell3.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_background(cell1, bg)
        set_cell_background(cell2, bg)
        set_cell_background(cell3, bg)
        set_cell_padding(cell1, top=90, bottom=90, left=130, right=130)
        set_cell_padding(cell2, top=90, bottom=90, left=130, right=130)
        set_cell_padding(cell3, top=90, bottom=90, left=130, right=130)

        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(c1)
        r1.font.name = 'Calibri'
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = TEXT_COLOR

        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(c2)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9)
        r2.font.color.rgb = TEXT_COLOR

        p3 = cell3.paragraphs[0]
        r3 = p3.add_run(c3)
        r3.font.name = 'Calibri'
        r3.font.size = Pt(9)
        r3.font.color.rgb = TEXT_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 4: ACCEPTANCE & KPI METRICS ---
    add_heading_1("4. Acceptance Criteria & KPI Success Metrics")
    add_body("The project deployment is validated against rigorous operational acceptance criteria:")

    add_bullet("Recognition Accuracy (>98%)", "Zero false-positive identity locks in workplace test scenarios through consecutive matching.")
    add_bullet("Stream Latency (<200ms)", "Live video rendering stays real-time with zero buffering lag across continuous multi-hour runs.")
    add_bullet("Working Time Accuracy (>99%)", "Net working hours match physical presence time within a 60-second grace tolerance.")
    add_bullet("False-Positive Phone Usage (<2%)", "Requires physical hand proximity and 2.0s duration, ignoring stationary phones on desks.")
    add_bullet("Automated Data Archival", "Structured JSON records auto-generated in output/known/ and output/unknown/ without manual intervention.")

    # Save
    workspace_path = "/Users/poornima/Video Analytics/Project_Plan_Document.docx"
    downloads_path = "/Users/poornima/Downloads/Project_Plan_Document.docx"
    
    doc.save(workspace_path)
    doc.save(downloads_path)
    print(f"Successfully generated Plan Word document at:\n1. {workspace_path}\n2. {downloads_path}")

if __name__ == "__main__":
    create_plan_document()
